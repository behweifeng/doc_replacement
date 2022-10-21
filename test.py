import pandas as pd
import re 
from docx import Document 
import sys

def docx_replace_regex(doc_obj, regex , replace): 

    paragraphs = list(doc.paragraphs)    

    for hp in doc.sections[0].header.paragraphs:
        paragraphs.append(hp)


    for ht in doc.sections[0].header.tables:
        for row in ht.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)

    for p in paragraphs: 
        if regex.search(p.text): 
            inline = p.runs 
            # Loop added to work with runs (strings with same style) 
            for i in range(len(inline)): 
                if regex.search(inline[i].text): 
                    text = regex.sub(replace, inline[i].text) 
                    inline[i].text = text 

try:
    docName = sys.argv[1]
    excelName = sys.argv[2]

    if docName.endswith(".docx") and excelName.endswith(".xlsx"):
        try:
            df = pd.read_excel(excelName)
            doc = Document(docName)

            for index, row in df.iterrows():
                doc = Document(docName)
                for col in df.columns:
                    regex1 = re.compile(str(col)) 
                    replace1 = str(row[col])
                    docx_replace_regex(doc, regex1 , replace1) 
                outputFileName = 'result'+str(index)+'.docx'
                doc.save(outputFileName) 
                print(outputFileName + " saved successfully.")

            print("Task completed successfully.")

        except:
            print("Error - one or more of the files cannot be opened.")
    else:
        print("Error - please check that files end with .docx and .xlsx")

except:
    print("Error - two arguments are needed in the command.")
    print("Example: python3 test.py <Name_Of_Doc>.docx <Name_Of_Excel>.xlsx")

